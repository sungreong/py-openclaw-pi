from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest


ROOT = Path(__file__).resolve().parents[1]
TOOL_PATH = ROOT / ".piagent" / "tools" / "word-report" / "tool.py"


def _load_module():
    spec = importlib.util.spec_from_file_location("piagent_word_report_test", TOOL_PATH)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_word_report_parses_markdown_table_cells():
    module = _load_module()

    assert module._table_cells("| **Title** | [Source](https://example.test) |") == [
        "Title",
        "Source (https://example.test)",
    ]
    assert module._is_table_separator("| --- | :---: |") is True
    assert module._plain_markdown('<span style="color:red;">RED</span>') == "RED"


def test_word_report_writes_explicit_table_geometry_and_header(tmp_path: Path):
    sys.path.append(str(ROOT / ".piagent" / "packages"))
    pytest.importorskip("docx")
    module = _load_module()
    output = tmp_path / "report.docx"

    stats = module._convert(
        "# Report\n\n| Risk | State |\n| --- | --- |\n| A | <span>RED</span> |\n",
        output,
    )

    from docx import Document

    document = Document(output)
    xml = document.tables[0]._tbl.xml
    assert stats["tables"] == 1
    assert 'w:type="dxa" w:w="8520"' in xml
    assert 'w:tblInd w:type="dxa" w:w="120"' in xml
    assert "w:tblHeader" in xml
    assert "<span" not in document.tables[0].cell(1, 1).text


def test_word_report_path_stays_inside_workspace(tmp_path: Path, monkeypatch):
    module = _load_module()
    monkeypatch.chdir(tmp_path)

    try:
        module._workspace_path("../outside.md")
    except ValueError as exc:
        assert "escapes workspace" in str(exc)
    else:
        raise AssertionError("path escape was accepted")


def test_word_report_blocks_sensitive_roots(tmp_path: Path, monkeypatch):
    module = _load_module()
    monkeypatch.chdir(tmp_path)

    try:
        module._workspace_path(".env")
    except ValueError as exc:
        assert "blocked" in str(exc)
    else:
        raise AssertionError("blocked path was accepted")
