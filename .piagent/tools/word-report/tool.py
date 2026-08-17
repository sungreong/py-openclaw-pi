from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any

from langchain.tools import tool


_MAX_MARKDOWN_BYTES = 2_000_000
_BLOCKED_ROOTS = {".git", ".env", ".openclaw", "secrets", "private", "node_modules"}


def _workspace_path(raw_path: str, *, must_exist: bool = False) -> Path:
    root = Path.cwd().resolve()
    candidate = Path(str(raw_path or "").strip())
    if not str(candidate):
        raise ValueError("path is required")
    resolved = (root / candidate).resolve() if not candidate.is_absolute() else candidate.resolve()
    try:
        relative = resolved.relative_to(root)
    except ValueError as exc:
        raise ValueError("path escapes workspace") from exc
    if relative.parts and relative.parts[0].lower() in _BLOCKED_ROOTS:
        raise ValueError("path is blocked by workspace policy")
    if must_exist and not resolved.is_file():
        raise ValueError(f"file not found: {raw_path}")
    return resolved


def _plain_markdown(value: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", str(value or ""))
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.replace("`", "").strip()


def _table_cells(line: str) -> list[str]:
    return [_plain_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def _set_default_font(document: Any) -> None:
    from docx.oxml.ns import qn

    style = document.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = __import__("docx").shared.Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_table_geometry(table: Any, column_count: int, total_width_twips: int = 8520) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips

    widths = [total_width_twips // column_count] * column_count
    widths[-1] += total_width_twips - sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.insert(0, table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(total_width_twips))

    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "120")

    table_grid = table._tbl.tblGrid
    grid_columns = list(table_grid.findall(qn("w:gridCol")))
    for index, width in enumerate(widths):
        if index < len(grid_columns):
            grid_columns[index].set(qn("w:w"), str(width))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            cell_width.type = "dxa"
            cell_width.w = width

    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_properties.append(header_marker)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _convert(markdown: str, output_path: Path, title: str = "") -> dict[str, int]:
    from docx import Document

    document = Document()
    _set_default_font(document)
    if title.strip():
        document.add_heading(_plain_markdown(title), level=0)

    lines = markdown.splitlines()
    stats = {"headings": 0, "paragraphs": 0, "tables": 0, "table_rows": 0, "list_items": 0}
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue

        if line.lstrip().startswith("|") and line.rstrip().endswith("|"):
            block: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                block.append(candidate)
                index += 1
            data_rows = [_table_cells(row) for row in block if not _is_table_separator(row)]
            if data_rows:
                column_count = max(len(row) for row in data_rows)
                table = document.add_table(rows=len(data_rows), cols=column_count)
                table.style = "Table Grid"
                for row_index, row in enumerate(data_rows):
                    for column_index, value in enumerate(row):
                        table.cell(row_index, column_index).text = value
                _set_table_geometry(table, column_count)
                stats["tables"] += 1
                stats["table_rows"] += len(data_rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            document.add_heading(_plain_markdown(heading.group(2)), level=min(len(heading.group(1)), 6))
            stats["headings"] += 1
            index += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            match = bullet or numbered
            document.add_paragraph(
                _plain_markdown(match.group(1)),
                style="List Bullet" if bullet else "List Number",
            )
            stats["list_items"] += 1
            index += 1
            continue

        document.add_paragraph(_plain_markdown(line))
        stats["paragraphs"] += 1
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return stats


@tool("word_report_create")
def word_report_create(markdown_path: str, docx_path: str, title: str = "") -> str:
    """Create a structured Word DOCX from a workspace Markdown file, including headings, lists, and tables."""
    try:
        source = _workspace_path(markdown_path, must_exist=True)
        destination = _workspace_path(docx_path)
        if destination.suffix.lower() != ".docx":
            raise ValueError("docx_path must end with .docx")
        if source.stat().st_size > _MAX_MARKDOWN_BYTES:
            raise ValueError("Markdown source exceeds the 2 MB safety limit")
        markdown = source.read_text(encoding="utf-8-sig")
        stats = _convert(markdown, destination, title)
    except ModuleNotFoundError as exc:
        if exc.name == "docx":
            return json.dumps(
                {
                    "status": "missing_dependency",
                    "package": "python-docx==1.2.0",
                    "next_action": "Use python_package_install with this exact package, then retry.",
                }
            )
        return json.dumps({"status": "error", "error": f"missing module: {exc.name}"})
    except (OSError, UnicodeError, ValueError) as exc:
        return json.dumps({"status": "error", "error": str(exc)})

    return json.dumps(
        {
            "status": "ok",
            "source": os.path.relpath(source, Path.cwd()),
            "output": os.path.relpath(destination, Path.cwd()),
            "bytes": destination.stat().st_size,
            **stats,
        },
        ensure_ascii=False,
    )


TOOLS = [word_report_create]
