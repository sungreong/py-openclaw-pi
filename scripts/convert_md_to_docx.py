import sys
from pathlib import Path
from docx import Document

def md_to_docx(md_path: str, docx_path: str):
    md_file = Path(md_path)
    docx_file = Path(docx_path)
    if not md_file.is_file():
        print(f"Markdown file not found: {md_path}")
        sys.exit(1)
    doc = Document()
    with md_file.open('r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            # Simple handling: headings start with '#'
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(line)
    doc.save(docx_file)
    print(f"Converted {md_path} to {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python convert_md_to_docx.py <input_md> <output_docx>')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
